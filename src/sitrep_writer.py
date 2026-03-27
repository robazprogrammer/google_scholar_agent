import os
from datetime import datetime

import pandas as pd
from docx import Document


def categorize_articles(articles):
    top = []
    maybe = []
    low = []

    for article in sorted(
        articles,
        key=lambda x: (
            x.get("core_aael", False),
            x.get("score", 0),
            x.get("title", "").lower(),
        ),
        reverse=True,
    ):
        score = article.get("score", 0)
        core = article.get("core_aael", False)

        if score >= 82 or (core and score >= 50):
            top.append(article)
        elif score >= 60:
            maybe.append(article)
        else:
            low.append(article)

    return top, maybe, low


def write_doc_section(doc, heading, articles, max_items=None):
    doc.add_heading(heading, level=1)

    if not articles:
        doc.add_paragraph("None.")
        return

    items = articles[:max_items] if max_items else articles

    for i, article in enumerate(items, start=1):
        title_line = article["title"]
        if article.get("core_aael", False):
            title_line += " [Core AAEL]"

        doc.add_heading(f"{i}. {title_line}", level=2)
        doc.add_paragraph(f"Score: {article.get('score', 0)}")
        doc.add_paragraph(f"Why it matters: {article.get('reason', '')}")
        doc.add_paragraph(f"Link: {article.get('url', '')}")

        summary = article.get("summary", "").strip()
        if summary:
            doc.add_paragraph(f"Snippet: {summary}")


def write_notebooklm_packet(path, top, maybe):
    with open(path, "w", encoding="utf-8") as f:
        f.write("AAEL Scholar Daily Packet\n")
        f.write("=" * 60 + "\n\n")

        f.write("TOP AAEL CANDIDATES\n")
        f.write("-" * 60 + "\n\n")

        for article in top[:12]:
            f.write(f"Title: {article['title']}\n")
            f.write(f"Score: {article.get('score', 0)}\n")
            f.write(f"Core AAEL: {article.get('core_aael', False)}\n")
            f.write(f"Why it matters: {article.get('reason', '')}\n")
            f.write(f"Link: {article.get('url', '')}\n")
            if article.get("summary", ""):
                f.write(f"Snippet: {article['summary']}\n")
            f.write("\n")

        f.write("\nPOSSIBLE AAEL CANDIDATES\n")
        f.write("-" * 60 + "\n\n")

        for article in maybe[:15]:
            f.write(f"Title: {article['title']}\n")
            f.write(f"Score: {article.get('score', 0)}\n")
            f.write(f"Core AAEL: {article.get('core_aael', False)}\n")
            f.write(f"Why it matters: {article.get('reason', '')}\n")
            f.write(f"Link: {article.get('url', '')}\n")
            if article.get("summary", ""):
                f.write(f"Snippet: {article['summary']}\n")
            f.write("\n")


def write_outputs(articles):
    os.makedirs("output", exist_ok=True)

    timestamp = datetime.now().strftime("%Y-%m-%d")
    doc_path = os.path.join("output", f"AAEL_SitRep_{timestamp}.docx")
    csv_path = os.path.join("output", f"AAEL_articles_{timestamp}.csv")
    txt_path = os.path.join("output", f"NotebookLM_packet_{timestamp}.txt")

    top, maybe, low = categorize_articles(articles)

    doc = Document()
    doc.add_heading(f"AAEL Scholar SitRep {timestamp}", level=0)
    doc.add_paragraph(f"Total deduplicated articles: {len(articles)}")
    doc.add_paragraph(f"Top/Core AAEL candidates: {len(top)}")
    doc.add_paragraph(f"Possible AAEL fit: {len(maybe)}")
    doc.add_paragraph(f"Weak/Low fit: {len(low)}")

    write_doc_section(doc, "Top AAEL Candidates", top, max_items=12)
    write_doc_section(doc, "Possible AAEL Candidates", maybe, max_items=15)

    doc.add_heading("Low Priority Items", level=1)
    doc.add_paragraph(
        "Lower-scoring items were kept in the CSV for tracking but omitted from the main briefing."
    )

    doc.save(doc_path)

    df = pd.DataFrame(articles)
    df = df.sort_values(
        by=["core_aael", "score", "title"],
        ascending=[False, False, True],
    )
    df.to_csv(csv_path, index=False)

    write_notebooklm_packet(txt_path, top, maybe)

    return [doc_path, csv_path, txt_path]
